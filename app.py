from flask import Flask, render_template,request,jsonify,redirect
import pandas as pd
from sqlalchemy import create_engine
import openai
from flask_mysqldb import MySQL
import mysql.connector
import docx
import spacy
import nltk


## Configure OpenAI API Key
openai.api_key = "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"

def create_connection():
    connection = mysql.connector.connect(
        host='localhost',
        user='root',
        password='',
        database='flask'
    )
    return connection


app = Flask(__name__)

@app.route("/")
def hello():

    return render_template('dashboard.html')

@app.route("/tenderProcess")
def harry():
    return render_template('tenderWritting.html')


@app.route('/uploadTender', methods=['GET', 'POST'])
def save_file():
    file_data = request.files['file']
    filename = request.form.get('filename')
    specifications = request.form.get('filename')
    print(specifications)
    file_data.save('uploadedfile/'+filename)
    

   
    # Open the document file
    doc = docx.Document('uploadedfile/'+filename)
  
    questions = []
    # Iterate through the tables in the document
    nlp = spacy.load('en_core_web_sm')
    for table in doc.tables:
    # Iterate through the rows in the table
        for row in table.rows:
            # Iterate through the cells in the row
            for cell in row.cells:
                # Iterate through the paragraphs in the cell
                for paragraph in cell.paragraphs:
                    # Check if the paragraph text ends with a question mark
                    doc = nlp(paragraph.text)
                    

                        # Loop through all the sentences in the paragraph
                    for sent in doc.sents:
                        words = nltk.word_tokenize(sent.text.strip())
                        # Check if the sentence ends with a question mark
                    
                        verb_in_second_position = (len(sent) > 2 and sent[0].pos_ == 'VERB')
                        if verb_in_second_position or (len(sent) > 2 and (words[0].lower() in ["what", "when", "where", "who", "whom", "whose", "which", "why", "how","if"] or words[-1] == "?")):
                            # Check if the sentence starts with a capital letter
                        
                            questions.append(sent.text.strip())
    data=[]
    connection = create_connection()
    cursor = connection.cursor()
    
    
    for question in questions:
        print(question)
        user= {"question":question}
        data.append(user)
        cursor.execute("SELECT question FROM question where question = '" + question +"'")
        result = cursor.fetchall()
        print("Result: ", result) 
        
        if result:
            ## If a match is found, return the answer 
            print("Result: ", result)
            
        else:
            cursor.execute("INSERT INTO question (question, answer,filename) VALUES (%s, %s, %s)", (question, 'answer', filename))

    ## Returning new answer as response
    cursor.execute("SELECT * FROM question")
    qdata = cursor.fetchall()
    connection.commit()
 
    #Closing the cursor
    cursor.close()
    return jsonify(qdata)
    
    #print(bullet_points);
@app.route("/deletequestion/<int:id>",methods=['GET','POST'])
def bootstrap(id):
    connection = create_connection()
    cursor = connection.cursor()
    cursor.execute("DELETE FROM questions where id =id")
    #return render_template('bootstrap.html')
    return redirect("/tenderProcess", code=302)

app.run(debug=True)
